from datetime import datetime, date
from io import BytesIO
from flask import Flask, render_template, request, redirect, url_for, flash, send_file, jsonify, abort, session
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy import func, case
from werkzeug.security import generate_password_hash, check_password_hash
from docx import Document
import os
from functools import wraps

app = Flask(__name__)
app.secret_key = "251098"

app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///inventory.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)

# MODELS________________________

class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(150), unique=True, nullable=False)
    password_hash = db.Column(db.String(255), nullable=False)
    role = db.Column(db.String(20), default="user")
    created_at = db.Column(db.DateTime, default=datetime.now)

    def set_password(self, password):
        self.password_hash = generate_password_hash(password)

    def check_password(self, password):
        return check_password_hash(self.password_hash, password)

class Product(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(255), unique=True, nullable=False, index=True)
    quantity = db.Column(db.Integer, nullable=False, default=0) #So luong ton hien tai
    created_at = db.Column(db.DateTime, default=datetime.now)

    material = db.Column(db.String(100))
    weight = db.Column(db.Float)
    draw_no = db.Column(db.String(50))

class InventoryTransaction(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    product_id = db.Column(db.Integer, db.ForeignKey('product.id'), nullable=False, index=True)
    type = db.Column(db.String(10), nullable=False)  # 'in' or 'out'
    quantity = db.Column(db.Integer, nullable=False)
    tx_date = db.Column(db.Date, nullable=False, index=True)
    created_at = db.Column(db.DateTime, default=datetime.now)

    product = db.relationship('Product', backref=db.backref('transactions', lazy=True))

#INIT DB
with app.app_context():
    db.create_all()

    #Tao admin mac dinh neu chua co
    if not User.query.filter_by(username="admin").first():
        admin = User(username="admin", role="admin")
        admin.set_password("quang@123")
        db.session.add(admin)
        db.session.commit()
        print("Tao user admin mac dinh")

#HELPERS________________________
def normalize_name(s: str) -> str:
    return ' '.join(s.strip().split())

def get_or_create_product(name: str):
    fixed = normalize_name(name)
    prod = Product.query.filter(func.lower(Product.name) == func.lower(fixed)).first()
    if not prod:
        prod = Product(name=fixed, quantity=0)
        db.session.add(prod)
        db.session.flush()
    return prod 

def stock_as_of(product_id: int, on_date: date) -> int:
    """Số lượng tồn của sản phẩm tính đến (và gồm) ngày on_date."""
    in_sum, out_sum = db.session.query(
        func.coalesce(func.sum(case((InventoryTransaction.type=='IN', InventoryTransaction.quantity), else_=0)), 0),
        func.coalesce(func.sum(case((InventoryTransaction.type=='OUT', InventoryTransaction.quantity), else_=0)), 0),
    ).filter(
        InventoryTransaction.product_id == product_id,
        InventoryTransaction.tx_date <= on_date
    ).one()
    return int(in_sum) - int(out_sum)

def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            flash("Vui lòng đăng nhập để tiếp tục.", "warning")
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

#ROUTES_________________________
@app.route('/')
@login_required
def index():
    # Vai so lieu thong ke
    total_products = Product.query.count()
    total_qty = db.session.query(func.coalesce(func.sum(Product.quantity), 0)).scalar() or 0

    # 5 giao dich gan nhat
    latest_txs = (InventoryTransaction.query.order_by(InventoryTransaction.created_at.desc())
                  .limit(5).all())

    # Top 5 sản phẩm tồn kho nhiều nhất
    top_products = Product.query.order_by(Product.quantity.desc()).limit(5).all()
    top_labels = [p.name for p in top_products]
    top_values = [p.quantity for p in top_products]
    return render_template('index.html', total_products=total_products,
                            total_qty=total_qty, latest_txs=latest_txs,
                            top_labels=top_labels, top_values=top_values)

@app.route('/stock-in', methods=['POST'])
def stock_in():
    name = request.form.get('name', '').strip()
    qty = request.form.get('quantity', '0')
    tx_date = request.form.get('tx_date', '')

    if not name:
        flash('Vui lòng nhập tên sản phẩm.', 'danger')
        return redirect(url_for('index'))
    
    try:
        qty = int(qty)
        if qty <= 0:
            raise ValueError()
    except:
        flash('Số luợng phải là một số nguyên dương.', 'danger')
        return redirect(url_for('index'))
    try:
        tx_date = datetime.strptime(tx_date, '%Y-%m-%d').date() if tx_date else date.today()
    except:
        flash('Ngày nhập không hợp lệ (đúng định dạng YYYY-MM-DD).', 'danger')    
        return redirect(url_for('index'))
    
    prod = get_or_create_product(name)
    prod.quantity += qty
    tx = InventoryTransaction(product_id=prod.id, type='IN', quantity=qty, tx_date=tx_date)
    db.session.add(tx)
    db.session.commit()

    flash(f'Nhập {qty} sản phẩm {prod.name} thành công.', 'success')
    return redirect(url_for('index'))

@app.route('/stock-out', methods=['POST'])
def stock_out():
    name = request.form.get('name', '').strip()
    qty = request.form.get('quantity', '0')
    tx_date = request.form.get('tx_date', '')

    if not name:
        flash('Vui lòng nhập tên sản phẩm.', 'danger')
        return redirect(url_for('index'))
    
    try:
        qty = int(qty)
        if qty <= 0:
            raise ValueError()
    except:
        flash('Số luợng phải là số nguyên dương.', 'danger')
        return redirect(url_for('index'))
    
    try:
        tx_date = datetime.strptime(tx_date, '%Y-%m-%d').date() if tx_date else date.today()
    except:
        flash('Ngày xuất không hợp lệ (đúng định dạng YYYY-MM-DD).', 'danger')    
        return redirect(url_for('index'))
    
    #Tìm sản phẩm
    prod = Product.query.filter(func.lower(Product.name) == func.lower(normalize_name(name))).first()
    if not prod:
        flash(f'Không tìm thấy sản phẩm {name}.', 'danger')
        return redirect(url_for('index'))
    
    # Khong cho phep ngay xuat truoc ngay nhap
    earliest_in = db.session.query(func.min(InventoryTransaction.tx_date)).filter(
        InventoryTransaction.product_id == prod.id,
        InventoryTransaction.type == 'IN'
    ).scalar()
    
    if earliest_in is None or tx_date < earliest_in:
        msg = 'Chưa có hàng nhập trước ngày này.' if earliest_in is None else f'Không thể xuất trước ngày nhập đầu tiên ({earliest_in}).'
        flash(msg, 'danger')
        return redirect(url_for('index'))
    
    # Kiểm tra tồn kho tính đến ngày xuất
    available_on_date = stock_as_of(prod.id, tx_date)
    if available_on_date < qty:
        flash(f'Vào ngày {tx_date} chỉ còn {available_on_date} sản phẩm {prod.name} trong kho, không đủ để xuất {qty} sản phẩm.', 'danger')
        return redirect(url_for('index'))
    
    # Nếu hợp lệ, thực hiện xuất hàng và cập nhật hàng tồn
    prod.quantity -= qty
    tx = InventoryTransaction(product_id=prod.id, type='OUT', quantity=qty, tx_date=tx_date)
    db.session.add(tx)
    db.session.commit()

    flash(f'Xuất {qty} sản phẩm {prod.name} thành công.', 'success')
    return redirect(url_for('index'))

@app.route('/products')
@login_required
def products():
    q = request.args.get('q', '').strip()
    query = Product.query
    if q:
        query = query.filter(Product.name.ilike(f'%{q}%'))
    items = query.order_by(Product.name.asc()).all()
    return render_template('products.html', items=items, q=q)

@app.route('/product/<int:product_id>')
def product_detail(product_id):
    prod = Product.query.get_or_404(product_id)
    # gom dữ liệu theo ngày cho biểu đồ (IN/OUT)
    rows = (db.session.query(
                InventoryTransaction.tx_date,
                func.sum(case((InventoryTransaction.type == 'IN', InventoryTransaction.quantity), else_=0)).label('qty_in'),
                func.sum(case((InventoryTransaction.type == 'OUT', InventoryTransaction.quantity), else_=0)).label('qty_out'),
            )
            .filter(InventoryTransaction.product_id == product_id)
            .group_by(InventoryTransaction.tx_date)
            .order_by(InventoryTransaction.tx_date.asc())
            .all())
    labels = [r.tx_date.strftime('%Y-%m-%d') for r in rows]
    qty_in = [int(r.qty_in or 0) for r in rows]
    qty_out = [int(r.qty_out or 0) for r in rows]
    return render_template('product_detail.html', prod=prod, labels=labels, qty_in=qty_in, qty_out=qty_out)

@app.route('/product/delete/<int:product_id>', methods=['POST'])
def delete_product(product_id):
    prod = Product.query.get_or_404(product_id)
    InventoryTransaction.query.filter_by(product_id=prod.id).delete()  # Xóa giao dịch liên quan
    
    db.session.delete(prod)
    db.session.commit()
    flash(f'Đã xóa sản phẩm {prod.name} và toàn bộ lịch sử giao dịch.', 'success')
    return redirect(url_for('products'))

@app.route("/product/edit/<int:prod_id>", methods=["GET", "POST"])
@login_required
def product_edit(prod_id):
    prod = Product.query.get_or_404(prod_id)

    if request.method == "POST":
        prod.material = request.form.get("material") or None
        try:
            prod.weight = float(request.form.get("weight")) if request.form.get('weight') else None
        except ValueError:
            flash("Khối lượng phải là số.", "danger")
            return redirect(url_for('product_edit', prod_id=prod.id))
        prod.draw_no = request.form.get("draw_no")
        db.session.commit()
        flash(f"Cập nhật thông tin sản phẩm thành công!", "success")
        return redirect(url_for("product_detail", product_id=prod.id))

    return render_template("product_edit.html", prod=prod)
    
@app.route('/transactions')
@login_required
def transactions():
    #Lọc theo tên, loại In/Out, từ ngày, đến ngày
    name = request.args.get('name', '').strip()
    ttype = request.args.get('type', '').strip().upper() # '', 'IN', 'OUT'
    start = request.args.get('start', '')
    end = request.args.get('end', '')

    query = (db.session.query(InventoryTransaction, Product)
             .join(Product, InventoryTransaction.product_id == Product.id))
    if name:
        query = query.filter(Product.name.ilike(f'%{name}%'))
    if ttype in ('IN', 'OUT'):
        query = query.filter(InventoryTransaction.type == ttype)
    if start:
        try:
            d = datetime.strptime(start, '%Y-%m-%d').date()
            query = query.filter(InventoryTransaction.tx_date >= d)
        except:
            flash('Ngày bắt đầu không hợp lệ (đúng định dạng YYYY-MM-DD).', 'warning')
        
    if end:
        try:
            d = datetime.strptime(end, '%Y-%m-%d').date()
            query = query.filter(InventoryTransaction.tx_date <= d)
        except:
            flash('Ngày kết thúc không hợp lệ (đúng định dạng YYYY-MM-DD).', 'warning')
    
    query = query.order_by(InventoryTransaction.tx_date.desc(), InventoryTransaction.created_at.desc())
    rows = query.all()
    return render_template('transactions.html', rows=rows, name=name, ttype=ttype, start=start, end=end)  

@app.route('/report', methods=['GET', 'POST'])
@login_required
def report():
    # GET: form chọn khoảng ngày; POST: tạo docx
    if request.method == 'GET':
        return render_template('reports.html')

    start = request.form.get('start', '')
    end = request.form.get('end', '')
    try:
        start_d = datetime.strptime(start, '%Y-%m-%d').date() if start else date.min
        end_d = datetime.strptime(end, '%Y-%m-%d').date() if end else date.max
    except:
        flash('Ngày không hợp lệ. Dùng định dạng YYYY-MM-DD.', 'danger')
        return redirect(url_for('report'))

    # Lấy tổng quan
    total_products = Product.query.count()
    total_qty = db.session.query(func.coalesce(func.sum(Product.quantity), 0)).scalar() or 0

    # Giao dịch trong khoảng
    txs = (db.session.query(InventoryTransaction, Product)
           .join(Product, InventoryTransaction.product_id == Product.id)
           .filter(InventoryTransaction.tx_date >= start_d,
                   InventoryTransaction.tx_date <= end_d)
           .order_by(InventoryTransaction.tx_date.asc())
           .all())

    # Tạo Word
    doc = Document()
    doc.add_heading('BÁO CÁO TỒN KHO', level=1)
    doc.add_paragraph(f'Khoảng thời gian: {start or "–"} đến {end or "–"}')
    doc.add_paragraph(f'Tổng số sản phẩm: {total_products}')
    doc.add_paragraph(f'Tổng số lượng tồn: {total_qty}')

    doc.add_heading('Giao dịch', level=2)
    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    hdr[0].text = 'Ngày'
    hdr[1].text = 'Sản phẩm'
    hdr[2].text = 'Loại'
    hdr[3].text = 'Số lượng'
    hdr[4].text = 'Tồn sau giao dịch (ước tính)'

    # để ước tính tồn sau giao dịch: duyệt theo thời gian
    stock_map = {p.id: p.quantity for p in Product.query.all()}
    # đảo ngược để suy ra dần không đáng tin trong thực tế; ở đây chỉ trình bày đơn giản.
    # Thực tế nên tính tiến theo thứ tự tăng dần từ tồn đầu kỳ. Để chính xác, ta tính tồn đầu kỳ:
    begin_stock = {}
    for p in Product.query.all():
        # tồn đầu kỳ = tồn hiện tại - (IN-OUT) trong [start_d, today]
        in_out_in_range = (db.session.query(
            func.sum(case((InventoryTransaction.type=='IN', InventoryTransaction.quantity), else_=0)).label('in_sum'),
            func.sum(case((InventoryTransaction.type=='OUT', InventoryTransaction.quantity), else_=0)).label('out_sum'),
        ).filter(InventoryTransaction.product_id==p.id,
                 InventoryTransaction.tx_date>=start_d).first())
        in_sum = int(in_out_in_range[0] or 0)
        out_sum = int(in_out_in_range[1] or 0)
        begin_stock[p.id] = p.quantity - (in_sum - out_sum)

    running = dict(begin_stock)
    # duyệt tx tăng dần ngày
    txs_sorted = sorted(txs, key=lambda x: (x[0].tx_date, x[0].created_at))
    for tx, prod in txs_sorted:
        if tx.type == 'IN':
            running[prod.id] = running.get(prod.id, 0) + tx.quantity
        else:
            running[prod.id] = running.get(prod.id, 0) - tx.quantity

        row = table.add_row().cells
        row[0].text = tx.tx_date.strftime('%Y-%m-%d')
        row[1].text = prod.name
        row[2].text = 'Nhập' if tx.type == 'IN' else 'Xuất'
        row[3].text = str(tx.quantity)
        row[4].text = str(running.get(prod.id, 0))

    # Xuất file
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f'BaoCao_TonKho_{start or "from"}_{end or "to"}.docx'
    return send_file(buf, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')  

# API nhỏ cho autocomplete tên sản phẩm
@app.route('/api/product-names')
def api_product_names():
    q = request.args.get('q', '').strip()
    query = Product.query
    if q:
        query = query.filter(Product.name.ilike(f'%{q}%'))
    names = [p.name for p in query.order_by(Product.name.asc()).limit(20).all()]
    return jsonify(names)

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')

        user = User.query.filter_by(username=username).first()
        if user and user.check_password(password):
            # Lưu thông tin đăng nhập vào session
            session['user_id'] = user.id
            session['username'] = user.username
            session['role'] = user.role
            flash('Đăng nhập thành công!', 'success')
            return redirect(url_for('index'))
        else:
            flash('Tên đăng nhập hoặc mật khẩu không đúng.', 'danger')
            return redirect(url_for('login'))
    return render_template('login.html', title='Đăng nhập')

@app.route('/logout')
def logout():
    session.clear()
    flash('Đã đăng xuất thành công.', 'info')
    return redirect(url_for('login'))

@app.route('/change-password', methods=['GET', 'POST'])
@login_required
def change_password():
    if 'user_id' not in session:
        flash('Bạn cần đăng nhập để thay đổi mật khẩu.', 'warning')
        return redirect(url_for('login'))

    user = User.query.get(session['user_id'])
    if not user:
        flash('Người dùng không tồn tại.', 'danger')
        return redirect(url_for('login'))

    if request.method == 'POST':
        old_password = request.form.get('old_password')
        new_password = request.form.get('new_password')
        confirm_password = request.form.get('confirm_password')

        if not user.check_password(old_password):
            flash('Mật khẩu cũ không đúng.', 'danger')
        elif new_password != confirm_password:
            flash('Mật khẩu mới và xác nhận không khớp.', 'danger')
        elif len(new_password) < 6:
            flash('Mật khẩu mới phải có ít nhất 6 ký tự.', 'warning')
        else:
            user.set_password(new_password)
            db.session.commit()
            flash('Đổi mật khẩu thành công.', 'success')
            return redirect(url_for('index'))
    return render_template('change_password.html', title='Đổi mật khẩu')

if __name__ == '__main__':
    # đảm bảo thư mục static/img tồn tại
    os.makedirs('static/img', exist_ok=True)
    app.run(host='0.0.0.0', port=5000)